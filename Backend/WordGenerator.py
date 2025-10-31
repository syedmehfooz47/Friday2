# -*- coding: utf-8 -*-
"""
Word Document Generator - Creates professional Word documents using Groq
 AI
Generates well-formatted, multi-page DOCX files
"""

import os
import time
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .logger import Logger
from .llm_handler import llm_handler

load_dotenv()


class WordGenerator:
    """Generate professional Word documents"""
    
    def __init__(self):
        self.data_folder = Path(__file__).parent.parent / "Data" / "GeneratedDocuments"
        self.data_folder.mkdir(parents=True, exist_ok=True)
    
    def generate_word(self, topic: str, pages: int = 10) -> tuple[str, str]:
        """
        Generate a professional Word document
        
        Args:
            topic: Topic for the document
            pages: Number of pages to generate (minimum 10)
            
        Returns:
            Tuple of (response message, file path)
        """
        username = os.getenv("Username", "Boss")
        
        if pages < 1:
            return f"Invalid page count, Boss. Please specify at least 1 page.", None
        
        Logger.log(f"Generating {pages}-page Word document on topic: '{topic}'", "WORD")
        
        try:
            # Generate content using Groq
            content_prompt = f"""You are a professional document writer.
Create a comprehensive, well-structured document about "{topic}".

The document should be approximately {pages * 400} words ({pages} pages worth of content).

Requirements:
- Start with an engaging introduction
- Use clear section headings (mark them with ### before the heading)
- Include detailed explanations and examples
- Write in a professional, informative tone
- End with a strong conclusion
- Format with proper paragraphs

Format headings like this:
### Section Heading
Content here...

### Another Heading
More content..."""
            
            messages = [{"role": "user", "content": content_prompt}]
            content = llm_handler.get_response(messages, max_tokens=8000, temperature=0.7)
            
            if content.startswith("Error"):
                return f"Failed to generate content: {content}", None
            
            # Create Word document
            doc = Document()
            
            # Title
            title = doc.add_heading(topic.upper(), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.color.rgb = RGBColor(44, 62, 80)
            title_run.font.size = Pt(24)
            
            doc.add_paragraph()
            
            # Parse and add content
            sections = content.split('###')
            
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                
                lines = section.split('\n', 1)
                
                if len(lines) == 2:
                    # It's a heading with content
                    heading_text = lines[0].strip()
                    body_text = lines[1].strip()
                    
                    # Add heading
                    heading = doc.add_heading(heading_text, level=2)
                    heading_run = heading.runs[0]
                    heading_run.font.color.rgb = RGBColor(52, 73, 94)
                    heading_run.font.size = Pt(16)
                    
                    # Add body paragraphs
                    paragraphs = body_text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            p = doc.add_paragraph(para.strip())
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            for run in p.runs:
                                run.font.size = Pt(12)
                                run.font.name = 'Calibri'
                else:
                    # Just body text
                    paragraphs = section.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            p = doc.add_paragraph(para.strip())
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            for run in p.runs:
                                run.font.size = Pt(12)
                                run.font.name = 'Calibri'
            
            # Save document
            safe_topic = "".join(c for c in topic if c.isalnum() or c in (' ', '_')).rstrip()[:50]
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            file_path = self.data_folder / f"{safe_topic.replace(' ', '_')}_{timestamp}.docx"
            
            doc.save(str(file_path))
            
            Logger.log(f"Word document generated successfully: {file_path}", "WORD")
            return f"I've generated a {pages}-page Word document on '{topic}' for you, Boss.", str(file_path)
        
        except Exception as e:
            Logger.log(f"Error generating Word document: {e}", "ERROR")
            return f"Failed to generate Word document: {e}", None


# Global instance
word_generator = WordGenerator()